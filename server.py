from __future__ import annotations

import io
import os
import re
import sqlite3
import time
import zipfile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import feedparser
from bs4 import BeautifulSoup
from docx import Document
from dateutil import parser as dateparser
from fastapi import Body, FastAPI, File, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles


APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(APP_DIR, "learnsite.db")


DEFAULT_FEEDS = [
    # IT / AI
    ("it", "Hugging Face Blog", "https://huggingface.co/blog/feed.xml"),
    ("it", "OpenAI News", "https://openai.com/news/rss.xml"),
    # Data / SQL
    ("it", "Postgres Weekly", "https://postgresweekly.com/rss/"),
    # Finance (偏英文；你可以后续替换为中文 RSS)
    ("finance", "CFA Institute (Insights)", "https://www.cfainstitute.org/en/research/rss"),
    ("finance", "Investopedia", "https://www.investopedia.com/feedbuilder/feed/getfeed?feedName=rss_articles"),
]


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def connect_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with connect_db() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS feeds (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              category TEXT NOT NULL,
              title TEXT NOT NULL,
              url TEXT NOT NULL UNIQUE,
              enabled INTEGER NOT NULL DEFAULT 1
            );
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS articles (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              feed_url TEXT NOT NULL,
              category TEXT NOT NULL,
              source_title TEXT NOT NULL,
              title TEXT NOT NULL,
              url TEXT NOT NULL UNIQUE,
              summary TEXT,
              published_at TEXT,
              fetched_at TEXT NOT NULL
            );
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS meta (
              key TEXT PRIMARY KEY,
              value TEXT NOT NULL
            );
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS knowledge_entries (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              source_name TEXT,
              source_type TEXT NOT NULL,
              category TEXT NOT NULL,
              title TEXT NOT NULL,
              summary TEXT,
              body TEXT NOT NULL,
              tags TEXT,
              source_path TEXT,
              created_at TEXT NOT NULL,
              updated_at TEXT NOT NULL
            );
            """
        )

        for category, title, url in DEFAULT_FEEDS:
            conn.execute(
                "INSERT OR IGNORE INTO feeds(category,title,url,enabled) VALUES(?,?,?,1)",
                (category, title, url),
            )


def set_meta(conn: sqlite3.Connection, key: str, value: str) -> None:
    conn.execute(
        "INSERT INTO meta(key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
        (key, value),
    )


def get_meta(conn: sqlite3.Connection, key: str) -> Optional[str]:
    row = conn.execute("SELECT value FROM meta WHERE key=?", (key,)).fetchone()
    return row["value"] if row else None


def parse_dt(entry: Any) -> Optional[str]:
    for k in ("published", "updated", "created"):
        v = getattr(entry, k, None)
        if not v:
            continue
        try:
            dt = dateparser.parse(v)
            if not dt:
                continue
            if not dt.tzinfo:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc).isoformat()
        except Exception:
            continue
    return None


def upsert_article(
    conn: sqlite3.Connection,
    *,
    feed_url: str,
    category: str,
    source_title: str,
    title: str,
    url: str,
    summary: Optional[str],
    published_at: Optional[str],
) -> int:
    before = conn.total_changes
    conn.execute(
        """
        INSERT OR IGNORE INTO articles(
          feed_url, category, source_title, title, url, summary, published_at, fetched_at
        ) VALUES (?,?,?,?,?,?,?,?)
        """,
        (feed_url, category, source_title, title, url, summary, published_at, utc_now_iso()),
    )
    after = conn.total_changes
    return 1 if after > before else 0


def fetch_once(limit_per_feed: int = 20) -> Dict[str, Any]:
    started = time.time()
    init_db()
    inserted = 0

    with connect_db() as conn:
        feeds = conn.execute("SELECT category,title,url FROM feeds WHERE enabled=1").fetchall()
        for f in feeds:
            category = f["category"]
            source_title = f["title"]
            feed_url = f["url"]

            parsed = feedparser.parse(feed_url)
            entries = getattr(parsed, "entries", [])[:limit_per_feed]
            for e in entries:
                title = (getattr(e, "title", "") or "").strip()
                url = (getattr(e, "link", "") or "").strip()
                if not title or not url:
                    continue
                summary = getattr(e, "summary", None) or getattr(e, "description", None) or None
                published_at = parse_dt(e)
                inserted += upsert_article(
                    conn,
                    feed_url=feed_url,
                    category=category,
                    source_title=source_title,
                    title=title,
                    url=url,
                    summary=summary,
                    published_at=published_at,
                )

        set_meta(conn, "last_fetch_at", utc_now_iso())

    return {
        "ok": True,
        "inserted": inserted,
        "elapsed_ms": int((time.time() - started) * 1000),
    }


def ensure_fresh(max_age_seconds: int = 60 * 60 * 6) -> Dict[str, Any]:
    init_db()
    with connect_db() as conn:
        last = get_meta(conn, "last_fetch_at")
    if not last:
        return fetch_once()
    try:
        dt = dateparser.parse(last)
        if not dt:
            return fetch_once()
        age = datetime.now(timezone.utc) - dt.astimezone(timezone.utc)
        if age.total_seconds() > max_age_seconds:
            return fetch_once()
        return {"ok": True, "skipped": True, "last_fetch_at": last}
    except Exception:
        return fetch_once()


def list_articles(category: Optional[str], limit: int) -> List[Dict[str, Any]]:
    init_db()
    with connect_db() as conn:
        if category and category != "all":
            rows = conn.execute(
                """
                SELECT category, source_title, title, url, summary, published_at
                FROM articles
                WHERE category=?
                ORDER BY COALESCE(published_at, fetched_at) DESC
                LIMIT ?
                """,
                (category, limit),
            ).fetchall()
        else:
            rows = conn.execute(
                """
                SELECT category, source_title, title, url, summary, published_at
                FROM articles
                ORDER BY COALESCE(published_at, fetched_at) DESC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()

    return [
        {
            "category": r["category"],
            "source": r["source_title"],
            "title": r["title"],
            "url": r["url"],
            "summary": (r["summary"] or "")[:300],
            "published_at": r["published_at"],
        }
        for r in rows
    ]


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def split_tags(text: str) -> List[str]:
    return [t.strip() for t in re.split(r"[，,\n\t]+", text or "") if t.strip()]


def guess_category(text: str) -> str:
    t = text.lower()
    if re.search(r"(sql|数据库|查询|表|窗口函数|join|select|insert|update|delete)", t, re.I):
        return "it"
    if re.search(r"(金融|报表|估值|现金流|资产负债|利润|债券|股票|宏观)", t, re.I):
        return "finance"
    return "general"


def extract_title_and_body(text: str, filename: str) -> tuple[str, str]:
    raw = normalize_whitespace(text.replace("\x00", " "))
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title = ""
    for line in lines[:20]:
        m = re.match(r"^#{1,3}\s+(.+)$", line)
        if m:
            title = m.group(1).strip()
            break
    if not title:
        title = filename.rsplit(".", 1)[0] if "." in filename else filename
    if not title:
        title = lines[0][:80] if lines else "未命名文档"
    body = "\n".join(lines) if lines else raw
    return normalize_whitespace(title), body.strip()


def extract_headings(text: str) -> List[str]:
    heads = []
    for line in text.splitlines():
        m = re.match(r"^#{1,6}\s+(.+)$", line.strip())
        if m:
            heads.append(m.group(1).strip())
    return heads


def text_from_docx_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def text_from_md_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def text_from_html_bytes(data: bytes) -> str:
    soup = BeautifulSoup(data, "lxml")
    return soup.get_text("\n")


def parse_document_bytes(data: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".docx"):
        return text_from_docx_bytes(data)
    if name.endswith(".md") or name.endswith(".markdown") or name.endswith(".txt"):
        return text_from_md_bytes(data)
    if name.endswith(".html") or name.endswith(".htm"):
        return text_from_html_bytes(data)
    return data.decode("utf-8", errors="ignore")


def summarize_text(text: str, max_len: int = 240) -> str:
    text = normalize_whitespace(text)
    if len(text) <= max_len:
        return text
    return text[: max_len - 1].rstrip() + "…"


def document_to_knowledge_entries(text: str, filename: str) -> List[Dict[str, Any]]:
    raw = text.strip()
    lines = [line.rstrip() for line in text.splitlines()]
    headings = extract_headings(text)
    title, body = extract_title_and_body(text, filename)
    category = guess_category(f"{title}\n{body}\n{filename}")
    tags = split_tags(",".join([
        "信息技术" if category == "it" else "金融" if category == "finance" else "通用",
        *headings[:3],
    ]))
    sections: List[Dict[str, Any]] = []

    if headings:
        current = []
        current_title = title
        for line in lines:
            m = re.match(r"^#{1,6}\s+(.+)$", line.strip())
            if m:
                if current:
                    section_body = "\n".join(current).strip()
                    sections.append({
                        "title": current_title,
                        "body": section_body,
                        "summary": summarize_text(section_body),
                    })
                    current = []
                current_title = m.group(1).strip()
            else:
                current.append(line)
        if current:
            section_body = "\n".join(current).strip()
            sections.append({
                "title": current_title,
                "body": section_body,
                "summary": summarize_text(section_body),
            })
    else:
        sections.append({"title": title, "body": body, "summary": summarize_text(body)})

    results: List[Dict[str, Any]] = []
    for idx, sec in enumerate(sections[:20]):
        sec_title = sec["title"] or title
        sec_body = sec["body"] or body or raw
        if not sec_body.strip():
            continue
        results.append(
            {
                "source_name": filename,
                "source_type": "document",
                "category": category,
                "title": sec_title[:120],
                "summary": sec["summary"],
                "body": sec_body.strip(),
                "tags": tags,
                "source_path": filename,
                "created_at": utc_now_iso(),
                "updated_at": utc_now_iso(),
            }
        )
    return results


def insert_knowledge_entries(conn: sqlite3.Connection, entries: List[Dict[str, Any]]) -> int:
    inserted = 0
    for e in entries:
        conn.execute(
            """
            INSERT INTO knowledge_entries(
              source_name, source_type, category, title, summary, body, tags, source_path, created_at, updated_at
            ) VALUES (?,?,?,?,?,?,?,?,?,?)
            """,
            (
                e.get("source_name"),
                e.get("source_type", "document"),
                e.get("category", "general"),
                e.get("title", "未命名文档"),
                e.get("summary", ""),
                e.get("body", ""),
                ",".join(e.get("tags", [])),
                e.get("source_path", ""),
                e.get("created_at", utc_now_iso()),
                e.get("updated_at", utc_now_iso()),
            ),
        )
        inserted += 1
    return inserted


app = FastAPI(title="LearnSite")


@app.on_event("startup")
def _startup() -> None:
    ensure_fresh()


@app.get("/api/health")
def health() -> Dict[str, Any]:
    return {"ok": True, "time": utc_now_iso()}


@app.post("/api/fetch")
def fetch() -> Dict[str, Any]:
    return fetch_once()


@app.get("/api/articles")
def articles(category: str = "all", limit: int = 30) -> JSONResponse:
    ensure_fresh()
    limit = max(1, min(int(limit), 100))
    data = list_articles(category, limit)
    return JSONResponse({"ok": True, "items": data})


@app.get("/api/pinned")
def pinned(count: int = 3) -> JSONResponse:
    ensure_fresh()
    count = max(1, min(int(count), 10))
    items = list_articles(None, 80)
    # “每日置顶自动更新”：用日期做稳定轮换（同一天不变）
    seed = datetime.now().strftime("%Y-%m-%d")
    h = 2166136261
    for ch in seed:
        h ^= ord(ch)
        h = (h * 16777619) & 0xFFFFFFFF
    x = h or 1
    arr = items[:]
    for i in range(len(arr) - 1, 0, -1):
        x = (1664525 * x + 1013904223) & 0xFFFFFFFF
        j = x % (i + 1)
        arr[i], arr[j] = arr[j], arr[i]
    return JSONResponse({"ok": True, "date": seed, "items": arr[:count]})


@app.post("/api/preview-doc")
async def preview_doc(file: UploadFile = File(...)) -> JSONResponse:
    filename = file.filename or "document.txt"
    data = await file.read()
    text = parse_document_bytes(data, filename)
    if not text.strip():
        return JSONResponse({"ok": False, "message": "文档内容为空"}, status_code=400)

    entries = document_to_knowledge_entries(text, filename)
    return JSONResponse({"ok": True, "filename": filename, "preview": entries})


@app.post("/api/import-doc/preview")
async def import_doc_preview(file: UploadFile = File(...)) -> JSONResponse:
    filename = file.filename or "document.txt"
    data = await file.read()
    text = parse_document_bytes(data, filename)
    if not text.strip():
        return JSONResponse({"ok": False, "message": "文档内容为空"}, status_code=400)
    preview = document_to_knowledge_entries(text, filename)
    return JSONResponse({"ok": True, "filename": filename, "preview": preview, "count": len(preview)})


@app.post("/api/import-doc/confirm")
async def import_doc_confirm(payload: Dict[str, Any]) -> JSONResponse:
    init_db()
    entries = payload.get("entries") or []
    if not isinstance(entries, list) or not entries:
        return JSONResponse({"ok": False, "message": "没有可入库条目"}, status_code=400)
    normalized = []
    for item in entries:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "source_name": item.get("source_name") or payload.get("filename") or "document",
                "source_type": item.get("source_type") or "document",
                "category": item.get("category") or "general",
                "title": item.get("title") or "未命名文档",
                "summary": item.get("summary") or "",
                "body": item.get("body") or "",
                "tags": item.get("tags") or [],
                "source_path": item.get("source_path") or payload.get("filename") or "",
                "created_at": item.get("created_at") or utc_now_iso(),
                "updated_at": item.get("updated_at") or utc_now_iso(),
            }
        )
    with connect_db() as conn:
        inserted = insert_knowledge_entries(conn, normalized)
        set_meta(conn, "last_doc_import_at", utc_now_iso())
    return JSONResponse({"ok": True, "inserted": inserted})


@app.get("/api/knowledge")
def knowledge(limit: int = 50) -> JSONResponse:


@app.get("/api/knowledge")
def knowledge(limit: int = 50) -> JSONResponse:
    init_db()
    limit = max(1, min(int(limit), 200))
    with connect_db() as conn:
        rows = conn.execute(
            """
            SELECT source_name, source_type, category, title, summary, body, tags, source_path, created_at, updated_at
            FROM knowledge_entries
            ORDER BY updated_at DESC, id DESC
            LIMIT ?
            """,
            (limit,),
        ).fetchall()
    items = [
        {
            "source_name": r["source_name"],
            "source_type": r["source_type"],
            "category": r["category"],
            "title": r["title"],
            "summary": r["summary"],
            "body": r["body"],
            "tags": [t for t in (r["tags"] or "").split(",") if t],
            "source_path": r["source_path"],
            "created_at": r["created_at"],
            "updated_at": r["updated_at"],
        }
        for r in rows
    ]
    return JSONResponse({"ok": True, "items": items})


app.mount("/", StaticFiles(directory=APP_DIR, html=True), name="static")


@app.get("/")
def root() -> FileResponse:
    return FileResponse(os.path.join(APP_DIR, "index.html"))

